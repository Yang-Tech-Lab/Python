"""
JurisOrchestrator Pro: v4.0 Strategic Legal Synthesis Engine
------------------------------------------------------------
A high-fidelity document orchestration node designed for resilient 
contract synthesis, featuring hardware-specific IP protection 
and automated audit trails.

Author: Yang Jiacheng (Yang-Tech-Lab)
Category: LegalTech / Systems Engineering
Date: May 2, 2026
"""

import logging
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Final, Optional, Dict

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. Industrial Infrastructure Configuration
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - [%(levelname)s] - %(message)s',
    handlers=[
        logging.FileHandler("Vault/Logs/contract_audit.log"),
        logging.StreamHandler()
    ]
)

@dataclass(frozen=True)
class EnterpriseClient:
    """High-fidelity client schema for contract orchestration."""
    name: str
    valuation_usd: str
    service_scope: str
    jurisdiction: str
    reference_id: str = field(default_factory=lambda: f"MSA-{datetime.now().strftime('%Y%m%d%H%M')}")

class JurisOrchestrator:
    def __init__(self, vault_dir: str = "Vault/Contracts"):
        self.vault_path: Final[Path] = Path(vault_dir)
        self.provider_id: Final[str] = "Yang Jiacheng (Yang-Tech-Lab)"
        
        # 2026 Industrial Aesthetic Palette
        self.brand_navy: Final[RGBColor] = RGBColor(28, 40, 51)  # Midnight Charcoal
        self.success_green: Final[RGBColor] = RGBColor(39, 174, 96) # Emerald Compliance
        
        self._bootstrap_vault()

    def _bootstrap_vault(self):
        """Provisions the secure local storage for synthesized assets."""
        self.vault_path.mkdir(parents=True, exist_ok=True)
        logging.info(f"🛠️ Orchestration Vault synchronized at: {self.vault_path.resolve()}")

    def _inject_typography(self, paragraph, size: int = 11, bold: bool = False, color: Optional[RGBColor] = None):
        """Orchestrates font styling with professional serif standards."""
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.name = 'Constantia'  # High-end Serif for Legal Clarity
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color

    def synthesize_contract(self, client: EnterpriseClient):
        """Orchestrates the synthesis of a high-fidelity Master Service Agreement."""
        logging.info(f"🚀 Initiating synthesis cycle for Node: {client.name}")
        doc = Document()
        
        # --- Section 1: Executive Header ---
        title = doc.add_heading('MASTER SERVICE AGREEMENT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata Block
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        meta.add_run(f"Contract ID: {client.reference_id}\n")
        meta.add_run(f"Effective Date: {datetime.now().strftime('%B %d, %Y')}\n")
        meta.add_run(f"Jurisdiction: {client.jurisdiction}")
        self._inject_typography(meta, size=9, color=colors.grey)

        # --- Section 2: Scope of Professional Deployment ---
        doc.add_heading('1. SCOPE OF ENGAGEMENT', level=1)
        sow = doc.add_paragraph("The Provider shall deliver specialized systems engineering including: ")
        sow_run = sow.add_run(client.service_scope)
        sow_run.bold = True
        sow_run.font.color.rgb = self.brand_navy

        # --- Section 3: Financial Consideration & Arbitrage ---
        doc.add_heading('2. REMUNERATION PROTOCOL', level=1)
        fee = doc.add_paragraph("The fixed consideration for this engagement is established at: ")
        fee_run = fee.add_run(f"${client.valuation_usd} USD")
        fee_run.bold = True
        fee_run.font.size = Pt(12)
        fee_run.font.color.rgb = self.success_green

        # --- Section 4: Hardware & Firmware IP (The Yang-Lab Standard) ---
        doc.add_heading('3. INTELLECTUAL PROPERTY & HARDWARE ASSETS', level=1)
        ip_p = doc.add_paragraph(
            "Upon full fiscal settlement, all deliverables—including KiCad PCB schematics, "
            "Gerber assets, BOM registries, and ESP-IDF firmware source code—shall be "
            "transferred to the Client. The Provider retains the right to utilize underlying "
            "automated libraries and generic hardware logic primitives."
        )
        self._inject_typography(ip_p)

        # --- Section 5: Execution Layer ---
        doc.add_page_break()
        doc.add_heading('EXECUTION OF AGREEMENT', level=1)
        
        table = doc.add_table(rows=2, cols=2)
        table.width = Inches(6)
        
        # Client Signature Cell
        c_cell = table.cell(0, 0)
        c_cell.text = f"FOR THE CLIENT ({client.name}):"
        
        # Provider Signature Cell
        p_cell = table.cell(0, 1)
        p_cell.text = f"FOR THE PROVIDER ({self.provider_id}):"
        
        # Signature Lines
        table.cell(1, 0).text = "\n\n__________________________\nAuthorized Signatory"
        table.cell(1, 1).text = f"\n\n__________________________\nSystem Architect"

        # --- Persistence Layer ---
        safe_name = client.name.replace(" ", "_").upper()
        target_file = self.vault_path / f"MSA_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        try:
            doc.save(str(target_file))
            logging.info(f"🏆 Strategic asset persisted: {target_file.name}")
        except Exception as e:
            logging.error(f"❌ Persistence Breach: {e}")

    def execute_batch_sequence(self, registry: List[EnterpriseClient]):
        """Orchestrates high-concurrency document synthesis."""
        print("\n" + "="*55)
        print("      YANG-TECH-LAB: JURIS-ORCHESTRATOR v4.0")
        print("="*55 + "\n")
        
        for node in registry:
            self.synthesize_contract(node)

if __name__ == "__main__":
    # Operational Registry: 2026.04.07
    nodes = [
        EnterpriseClient(
            name="Tesla Giga Guangzhou", 
            valuation_usd="45,000", 
            service_scope="Autonomous Hardware Validation Framework", 
            jurisdiction="Guangdong, PRC"
        ),
        EnterpriseClient(
            name="SpaceX", 
            valuation_usd="18,500", 
            service_scope="ESP32-S3 Satellite Telemetry Node", 
            jurisdiction="Texas, USA"
        )
    ]

    engine = JurisOrchestrator()
    engine.execute_batch_sequence(nodes)
